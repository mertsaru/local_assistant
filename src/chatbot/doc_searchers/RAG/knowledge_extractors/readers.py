import os
from typing import LiteralString


from docx import Document
import pdfplumber
from tabulate import tabulate


def _pdf_reader(file_path: LiteralString) -> LiteralString:
    """extracts the content with tables in markdown format

    Args:
        file_path (LiteralString): Full or relative path of the file

    Returns:
        LiteralString: content in markdown format
    """

    content = ""
    table_number = 0
    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages, 1):
            content += f"\n\n## Page {page_number}\n"

            # Extract text
            text = page.extract_text()
            if text:
                content += f"\n{text}\n"

            # Extract tables
            tables = page.extract_tables()
            for table in tables:
                if table:
                    table_number += 1
                    content += f"\ntable-{table_number}\n"
                    table = tabulate(table, headers="firstrow")
                    content += table + "\n"

    return content


def _docx_reader(file_path) -> LiteralString:

    doc = Document(file_path)
    content = ""

    # Extract text paragraphs
    for paragraph_number, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if text:
            content += f"\nparagraph-{paragraph_number}\n" + text + "\n"

    # Extract tables
    for table_number, table in enumerate(doc.tables, 1):
        content += f"\ntable-{table_number}\n"
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            content += " | ".join(row_text) + "\n"

    return content


def text_reader(file_path: LiteralString) -> tuple[LiteralString, LiteralString]:
    """extracts the content and the file name. The following types are accepted:

    - PDF
    - DOCX
    - TXT
    - Markdown

    Args:
        file_path (LiteralString): Full or relative path of the file

    Returns:
        tuple[LiteralString,LiteralString]: (file name, file content)
    """

    file_name = os.path.basename(file_path)

    file_extension = file_name.split(".")[-1].lower()

    if file_extension == "pdf":
        content = _pdf_reader(file_path)
    elif file_extension == "docx":
        content = _docx_reader(file_path)
    else:
        raise TypeError(f"file format {file_extension} is not supported")

    return file_name, content
